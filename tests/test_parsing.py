"""Parser tests: format dispatch, docx table -> markdown, document-order interleaving."""

from pathlib import Path

import pytest
from docx import Document

from parsing import parse_document
from parsing import _repair_rupee

SAMPLES = Path(__file__).parent.parent / "data" / "samples"


def test_unsupported_extension_raises(tmp_path: Path) -> None:
    bogus = tmp_path / "file.txt"
    bogus.write_text("hello")
    with pytest.raises(ValueError, match="Unsupported file type"):
        parse_document(bogus)


def test_docx_table_becomes_markdown(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Before table")
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Milestone"
    table.rows[0].cells[1].text = "Due Date"
    table.rows[0].cells[2].text = "Payment"
    table.rows[1].cells[0].text = "Milestone 1"
    table.rows[1].cells[1].text = "20 Sep 2024"
    table.rows[1].cells[2].text = "₹8,00,000"
    doc.add_paragraph("After table")
    f = tmp_path / "t.docx"
    doc.save(str(f))

    text = parse_document(f)
    assert "| Milestone | Due Date | Payment |" in text
    assert "| Milestone 1 | 20 Sep 2024 | ₹8,00,000 |" in text
    # interleaving: paragraph order preserved around the table
    assert text.index("Before table") < text.index("| Milestone |") < text.index("After table")


def test_multiline_cell_flattened(tmp_path: Path) -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "Ramesh Iyer, CTO\nramesh.iyer@novabridge.in"
    f = tmp_path / "t.docx"
    doc.save(str(f))
    assert "| Ramesh Iyer, CTO; ramesh.iyer@novabridge.in |" in parse_document(f)


def test_real_sow_milestone_row_association() -> None:
    """The SOW milestone table must keep milestone <-> amount <-> date on one line."""
    text = parse_document(SAMPLES / "doc_003_statement_of_work.docx")
    milestone_line = next(line for line in text.splitlines() if "Milestone 1" in line)
    assert "20 Sep 2024" in milestone_line
    assert "₹8,00,000" in milestone_line


def test_real_pdf_extracts_key_ids() -> None:
    text = parse_document(SAMPLES / "doc_004_invoice.pdf")
    assert "ENT-INV-2024-0043" in text
    assert "ENT-SOW-2024-019" in text


def test_rupee_repair_with_indian_context() -> None:
    text = "GSTIN: 09ABC | Subtotal: I12,00,000 | GST (18%): I 2,16,000"
    out = _repair_rupee(text)
    assert "₹12,00,000" in out
    assert "₹2,16,000" in out          # trailing space after I is consumed
    assert "I12" not in out and "I 2" not in out


def test_rupee_repair_skipped_without_indian_context() -> None:
    text = "Annual fee of $120,000 USD; Section I 5 applies."
    assert _repair_rupee(text) == text  # no GST/IFSC/crore -> untouched


def test_dollar_amounts_never_corrupted() -> None:
    text = "GSTIN present. Fee: $120,000 USD/year. Total: I40,00,000"
    out = _repair_rupee(text)
    assert "$120,000" in out            # $ is never matched
    assert "₹40,00,000" in out


def test_real_invoice_amounts_repaired() -> None:
    text = parse_document(SAMPLES / "doc_004_invoice.pdf")
    assert "₹14,16,000" in text
    assert "I14,16,000" not in text


def test_real_contract_has_no_rupee_introduced() -> None:
    text = parse_document(SAMPLES / "doc_001_vendor_contract.docx")
    assert "$120,000" in text
    assert "₹" not in text              # gate stays closed on the dollar contract
